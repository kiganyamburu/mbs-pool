from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
import re

# Read the markdown file
with open("Homework5_Answers.docx.md", "r", encoding="utf-8") as f:
    content = f.read()

# Create Word document
doc = Document()

# Set default font
style = doc.styles["Normal"]
font = style.font
font.name = "Calibri"
font.size = Pt(11)

# Process markdown content
lines = content.split("\n")
i = 0
in_table = False
table_rows = []
in_code_block = False

while i < len(lines):
    line = lines[i]

    # Skip empty lines
    if not line.strip():
        if in_table and table_rows:
            # End of table, create it
            if len(table_rows) > 0:
                # Filter out separator rows
                data_rows = [
                    r
                    for r in table_rows
                    if not all(
                        c.strip().replace("-", "").replace("|", "") == "" for c in r
                    )
                ]
                if data_rows:
                    num_cols = len(data_rows[0])
                    table = doc.add_table(rows=len(data_rows), cols=num_cols)
                    table.style = "Table Grid"
                    for row_idx, row_data in enumerate(data_rows):
                        for col_idx, cell_data in enumerate(row_data):
                            if col_idx < len(table.rows[row_idx].cells):
                                table.rows[row_idx].cells[
                                    col_idx
                                ].text = cell_data.strip()
            table_rows = []
            in_table = False
        i += 1
        continue

    # Handle code blocks
    if line.strip().startswith("```"):
        in_code_block = not in_code_block
        i += 1
        continue

    if in_code_block:
        p = doc.add_paragraph()
        run = p.add_run(line)
        run.font.name = "Consolas"
        run.font.size = Pt(10)
        i += 1
        continue

    # Handle tables
    if "|" in line and line.strip().startswith("|"):
        in_table = True
        # Parse table row
        cells = [
            c.strip() for c in line.split("|")[1:-1]
        ]  # Remove first and last empty
        # Skip separator rows
        if not all(c.replace("-", "").replace(":", "").strip() == "" for c in cells):
            table_rows.append(cells)
        i += 1
        continue

    # Handle headers
    if line.startswith("# "):
        doc.add_heading(line[2:].strip(), level=0)
        i += 1
        continue
    elif line.startswith("## "):
        doc.add_heading(line[3:].strip(), level=1)
        i += 1
        continue
    elif line.startswith("### "):
        doc.add_heading(line[4:].strip(), level=2)
        i += 1
        continue
    elif line.startswith("#### "):
        doc.add_heading(line[5:].strip(), level=3)
        i += 1
        continue

    # Handle horizontal rules
    if line.strip() == "---":
        doc.add_paragraph("─" * 50)
        i += 1
        continue

    # Handle bullet points
    if line.strip().startswith("- ") or line.strip().startswith("* "):
        text = line.strip()[2:]
        # Clean up markdown formatting
        text = re.sub(r"\*\*(.+?)\*\*", r"\1", text)  # Bold
        text = re.sub(r"\*(.+?)\*", r"\1", text)  # Italic
        text = re.sub(r"`(.+?)`", r"\1", text)  # Code
        p = doc.add_paragraph(text, style="List Bullet")
        i += 1
        continue

    # Handle numbered lists
    match = re.match(r"^(\d+)\.\s+(.+)", line.strip())
    if match:
        text = match.group(2)
        text = re.sub(r"\*\*(.+?)\*\*", r"\1", text)
        text = re.sub(r"\*(.+?)\*", r"\1", text)
        text = re.sub(r"`(.+?)`", r"\1", text)
        p = doc.add_paragraph(text, style="List Number")
        i += 1
        continue

    # Handle regular text with markdown cleanup
    text = line.strip()
    if text:
        # Clean up markdown formatting for display
        text = re.sub(r"\*\*(.+?)\*\*", r"\1", text)  # Bold markers
        text = re.sub(r"\*(.+?)\*", r"\1", text)  # Italic markers
        text = re.sub(r"`(.+?)`", r"\1", text)  # Code markers
        text = re.sub(r"\[([^\]]+)\]\([^\)]+\)", r"\1", text)  # Links
        text = re.sub(r"\$\$(.+?)\$\$", r"\1", text)  # Math blocks
        text = re.sub(r"\$(.+?)\$", r"\1", text)  # Inline math

        if text.strip():
            doc.add_paragraph(text)

    i += 1

# Handle any remaining table
if in_table and table_rows:
    data_rows = [
        r
        for r in table_rows
        if not all(c.strip().replace("-", "").replace("|", "") == "" for c in r)
    ]
    if data_rows:
        num_cols = len(data_rows[0])
        table = doc.add_table(rows=len(data_rows), cols=num_cols)
        table.style = "Table Grid"
        for row_idx, row_data in enumerate(data_rows):
            for col_idx, cell_data in enumerate(row_data):
                if col_idx < len(table.rows[row_idx].cells):
                    table.rows[row_idx].cells[col_idx].text = cell_data.strip()

# Save the document
doc.save("Homework5_Answers_v2.docx")
print("Word document created: Homework5_Answers_v2.docx")
