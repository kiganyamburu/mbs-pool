"""Convert Spring_2026_Midterm_Answers.md to Word document"""

from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import re
import os

doc = Document()

# Read markdown file
with open("Spring_2026_Midterm_Answers.md", "r", encoding="utf-8") as f:
    content = f.read()

lines = content.split("\n")
i = 0
while i < len(lines):
    line = lines[i]

    # Skip empty lines
    if not line.strip():
        i += 1
        continue

    # Handle images
    img_match = re.match(r"!\[([^\]]*)\]\(([^)]+)\)", line)
    if img_match:
        img_path = img_match.group(2)
        if os.path.exists(img_path):
            doc.add_picture(img_path, width=Inches(6))
        i += 1
        continue

    # Handle headers
    if line.startswith("# "):
        p = doc.add_heading(line[2:], level=0)
    elif line.startswith("## "):
        doc.add_heading(line[3:], level=1)
    elif line.startswith("### "):
        doc.add_heading(line[4:], level=2)
    elif line.startswith("---"):
        doc.add_paragraph("_" * 50)
    elif line.startswith("| "):
        # Handle table - collect all table rows
        table_lines = []
        while i < len(lines) and lines[i].startswith("|"):
            table_lines.append(lines[i])
            i += 1
        i -= 1  # Back up one since we'll increment at end

        # Parse table
        if len(table_lines) > 1:
            # Get headers
            headers = [cell.strip() for cell in table_lines[0].split("|")[1:-1]]
            # Skip separator line
            data_start = (
                2
                if table_lines[1].replace("|", "").replace("-", "").replace(" ", "")
                == ""
                else 1
            )

            # Create table
            num_cols = len(headers)
            num_rows = len(table_lines) - data_start + 1
            table = doc.add_table(rows=num_rows, cols=num_cols)
            table.style = "Table Grid"

            # Add headers
            for j, header in enumerate(headers):
                table.rows[0].cells[j].text = header

            # Add data rows
            for row_idx, table_line in enumerate(table_lines[data_start:], 1):
                cells = [cell.strip() for cell in table_line.split("|")[1:-1]]
                for j, cell in enumerate(cells):
                    if j < num_cols:
                        table.rows[row_idx].cells[j].text = cell
    elif line.startswith("```"):
        # Code block - collect until closing ```
        code_lines = []
        i += 1
        while i < len(lines) and not lines[i].startswith("```"):
            code_lines.append(lines[i])
            i += 1
        code_text = "\n".join(code_lines)
        p = doc.add_paragraph()
        run = p.add_run(code_text)
        run.font.name = "Courier New"
        run.font.size = Pt(9)
    elif line.startswith("- ") or line.startswith("* "):
        # Bullet point
        doc.add_paragraph(line[2:], style="List Bullet")
    elif line.startswith("$$"):
        # Math block
        p = doc.add_paragraph(line)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    else:
        # Regular paragraph - handle bold and other formatting
        text = line
        # Remove markdown bold markers for cleaner display
        text = re.sub(r"\*\*([^*]+)\*\*", r"\1", text)
        text = re.sub(r"\$([^$]+)\$", r"\1", text)  # Remove inline math markers
        doc.add_paragraph(text)

    i += 1

doc.save("Spring_2026_Midterm_Answers_v2.docx")
print("Created Spring_2026_Midterm_Answers_v2.docx")
